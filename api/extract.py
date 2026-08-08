import csv
import io
import os

import pymupdf
import openpyxl
import pytesseract
from docx import Document
from PIL import Image

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}


def extension_of(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def extract_text(filename: str, data: bytes) -> str:
    ext = extension_of(filename)
    if ext not in (".txt", ".json", ".csv", ".pdf", ".docx", ".xlsx") and ext not in IMAGE_EXTS:
        raise ValueError(f"unsupported file type: {ext}")
    try:
        if ext in (".txt", ".json"):
            return data.decode("utf-8", errors="replace")
        if ext == ".csv":
            return _extract_csv(data)
        if ext == ".pdf":
            return _extract_pdf(data)
        if ext == ".docx":
            return _extract_docx(data)
        if ext == ".xlsx":
            return _extract_xlsx(data)
        return _extract_image(data)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"could not read {ext} file: {e}") from e


def _extract_csv(data: bytes) -> str:
    reader = csv.reader(io.StringIO(data.decode("utf-8", errors="replace")))
    return "\n".join("\n\n".join(row) for row in reader)


def _extract_pdf(data: bytes) -> str:
    parts = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text())
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("no text layer found - this PDF may be scanned (OCR not supported yet)")
    return text


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append("\n\n".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _extract_xlsx(data: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append("\n\n".join(cells))
    wb.close()
    return "\n".join(lines)


def _extract_image(data: bytes) -> str:
    text, _ = extract_image_with_boxes(data)
    return text


def extract_image_with_boxes(data: bytes):
    img = Image.open(io.BytesIO(data))
    gray = img.convert("L")
    bw = gray.point(lambda x: 0 if x < 180 else 255, "1").convert("L")
    ocr = pytesseract.image_to_data(bw, config="--psm 6 --oem 1", output_type=pytesseract.Output.DICT)
    text = ""
    boxes = []
    prev_line_key = None
    for i in range(len(ocr["text"])):
        word = ocr["text"][i]
        if not word.strip():
            continue
        line_key = (ocr["block_num"][i], ocr["par_num"][i], ocr["line_num"][i])
        if prev_line_key is not None:
            text += "\n" if line_key != prev_line_key else " "
        start = len(text)
        text += word
        boxes.append({
            "start": start, "end": len(text),
            "left": ocr["left"][i], "top": ocr["top"][i],
            "width": ocr["width"][i], "height": ocr["height"][i],
        })
        prev_line_key = line_key
    if not text:
        raise ValueError("no text found in image")
    return text, boxes
