import csv
import io

import pymupdf
import openpyxl
from docx import Document
from PIL import Image, ImageDraw

from .extract import extension_of, extract_image_with_boxes

TYPES = {
    ".txt": "text/plain; charset=utf-8",
    ".json": "application/json",
    ".csv": "text/csv; charset=utf-8",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
    ".bmp": "image/bmp",
}


def _is_specific(original):
    t = original.strip()
    return len(t) >= 6 or " " in t


def build_redaction_pairs_from_dicts(text, results, split_tokens=False):
    counts, order = {}, []
    for r in sorted(results, key=lambda r: (r["end"] - r["start"], r.get("score", 0)), reverse=True):
        start, end = r["start"], r["end"]
        if start < 0 or end > len(text) or start >= end:
            continue
        original = text[start:end]
        fragments = original.split("\n\n") if "\n\n" in original else [original]
        for fragment in fragments:
            if not fragment.strip() or len(fragment.strip()) < 3:
                continue
            entity_type = "EMAIL_ADDRESS" if "@" in fragment and " " not in fragment.strip() else r["entity_type"]
            replacement = f"<{entity_type}>"
            key = (fragment, replacement)
            if key not in counts:
                counts[key] = 0
                order.append(key)
            counts[key] += 1
            if split_tokens:
                for token in fragment.split():
                    if not token.strip():
                        continue
                    tkey = (token, replacement)
                    if tkey not in counts:
                        counts[tkey] = 0
                        order.append(tkey)
                    counts[tkey] += 1
    pairs = []
    for original, replacement in order:
        budget = None if _is_specific(original) else counts[(original, replacement)]
        pairs.append((original, replacement, budget))
    pairs.sort(key=lambda p: len(p[0]), reverse=True)
    return pairs


def _apply_replacements(s, pairs, remaining):
    for original, replacement, _ in pairs:
        budget = remaining.get(original, 0)
        if budget == 0:
            continue
        if budget is None:
            s = s.replace(original, replacement)
            continue
        occurrences = s.count(original)
        use = min(occurrences, budget)
        if use <= 0:
            continue
        s = s.replace(original, replacement, use)
        remaining[original] -= use
    return s


def redact_file(filename, data, text, results):
    ext = extension_of(filename)
    if ext == ".pdf":
        out = _redact_pdf(data, text, results)
    else:
        pairs = build_redaction_pairs_from_dicts(text, results, split_tokens=ext == ".csv")
        remaining = {original: count for original, _, count in pairs}
        if ext in (".txt", ".json"):
            out = _apply_replacements(data.decode("utf-8", errors="replace"), pairs, remaining).encode("utf-8")
        elif ext == ".csv":
            out = _redact_csv(data, pairs, remaining)
        elif ext == ".docx":
            out = _redact_docx(data, pairs, remaining)
        elif ext == ".xlsx":
            out = _redact_xlsx(data, pairs, remaining)
        else:
            raise ValueError(f"unsupported file type: {ext}")
    return out, TYPES[ext]


def redact_image_file(data, result_dicts):
    img = Image.open(io.BytesIO(data))
    fmt = img.format or "PNG"
    if img.mode not in ("RGB", "RGBA"):
        img = img.convert("RGB")
    _, boxes = extract_image_with_boxes(data)
    draw = ImageDraw.Draw(img)
    for r in result_dicts:
        for box in boxes:
            if box["start"] < r["end"] and box["end"] > r["start"]:
                x0, y0 = box["left"], box["top"]
                x1, y1 = x0 + box["width"], y0 + box["height"]
                draw.rectangle([x0, y0, x1, y1], fill=(0, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    ext = f".{fmt.lower()}"
    if ext == ".jpeg":
        ext = ".jpg"
    return buf.getvalue(), TYPES.get(ext, f"image/{fmt.lower()}")


def _redact_csv(data, pairs, remaining):
    rows = list(csv.reader(io.StringIO(data.decode("utf-8", errors="replace"))))
    for row in rows:
        for i, field in enumerate(row):
            row[i] = _apply_replacements(field, pairs, remaining)
    buf = io.StringIO()
    csv.writer(buf).writerows(rows)
    return buf.getvalue().encode("utf-8")


def _redact_paragraph(p, pairs, remaining):
    for run in p.runs:
        new = _apply_replacements(run.text, pairs, remaining)
        if new != run.text:
            run.text = new


def _redact_docx(data, pairs, remaining):
    doc = Document(io.BytesIO(data))
    for p in doc.paragraphs:
        _redact_paragraph(p, pairs, remaining)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _redact_paragraph(p, pairs, remaining)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _redact_xlsx(data, pairs, remaining):
    wb = openpyxl.load_workbook(io.BytesIO(data))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    new = _apply_replacements(cell.value, pairs, remaining)
                    if new != cell.value:
                        cell.value = new
    out = io.BytesIO()
    wb.save(out)
    wb.close()
    return out.getvalue()


def _redact_pdf(data, text, results):
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        page_texts = [doc[i].get_text() for i in range(len(doc))]
        page_offsets = []
        pos = 0
        for pt in page_texts:
            page_offsets.append(pos)
            pos += len(pt) + 1

        page_spans = {}
        for r in results:
            start, end = r["start"], r["end"]
            page_num = None
            for i, page_start in enumerate(page_offsets):
                page_end = page_start + len(page_texts[i])
                if page_start <= start < page_end + 1:
                    page_num = i
                    break
            if page_num is None:
                continue
            local_start = max(0, start - page_offsets[page_num])
            local_end = min(len(page_texts[page_num]), end - page_offsets[page_num])
            if local_start >= local_end:
                continue
            matched = page_texts[page_num][local_start:local_end]
            if not matched.strip():
                continue
            page_spans.setdefault(page_num, []).append((local_start, matched))

        for page_num, spans in page_spans.items():
            page = doc[page_num]
            page_text = page_texts[page_num]
            redacted_globally = set()
            for local_start, matched in spans:
                rects = page.search_for(matched)
                if not rects:
                    continue
                if _is_specific(matched):
                    if matched in redacted_globally:
                        continue
                    redacted_globally.add(matched)
                    for rect in rects:
                        page.add_redact_annot(rect, text="", fill=(0, 0, 0))
                else:
                    occurrence_index = page_text[:local_start].count(matched)
                    if occurrence_index < len(rects):
                        page.add_redact_annot(rects[occurrence_index], text="", fill=(0, 0, 0))
                    else:
                        for rect in rects:
                            page.add_redact_annot(rect, text="", fill=(0, 0, 0))
            page.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE)
        out = io.BytesIO()
        doc.save(out, garbage=4, deflate=True, clean=True)
    return out.getvalue()
